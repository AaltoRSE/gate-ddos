import tempfile
import unittest
import sys
import re
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from docx import Document

from gate_ddos.docx import DocxPipeline
from gate_ddos.docx.styles import ensure_required_styles
from gate_ddos.docx.html import postprocess_html
from gate_ddos.docx.markdown import normalize_newlines
from gate_ddos.models import SectionRecord, TemplateSyntax
from gate_ddos.section_store import SectionStore


class _DocxPipelineMixin:
    def _docx_xml(self, docx_path: Path) -> tuple[str, str]:
        with zipfile.ZipFile(docx_path) as zf:
            rels = zf.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
            docxml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        return rels, docxml

    def _process_docx(self, template_path: Path, output_path: Path, *, store=None, generate=None):
        DocxPipeline(
            "sys",
            TemplateSyntax(),
            store or SectionStore(),
            generate or (lambda *_: "unused"),
        ).process(template_path, output_path)


class DocxPipelineTests(_DocxPipelineMixin, unittest.TestCase):

    def test_multiline_placeholder_spanning_paragraphs(self):
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ SYSTEM_DESCRIPTION || Describe the system")
            template.add_paragraph("from a non-technical point of view. }}")
            template.save(str(template_path))

            def fake_generate(system_prompt, prompt):
                self.assertIn("non-technical", prompt)
                return "Generated description"

            self._process_docx(template_path, output_path, generate=fake_generate)

            result = Document(str(output_path))
            joined = "\n".join(p.text for p in result.paragraphs)
            self.assertIn("Generated description", joined)
            self.assertNotIn("{{", joined)

    def test_markdown_link_becomes_clickable_hyperlink(self):
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ CONTENT || p }}")
            template.save(str(template_path))

            store = SectionStore({
                "CONTENT": SectionRecord(
                    prompt="p",
                    output="See [Google](https://www.google.com) for details.",
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            rels, docxml = self._docx_xml(output_path)
            self.assertIn("https://www.google.com", rels)
            self.assertIn("<w:hyperlink", docxml)

    def test_bare_url_stays_as_plain_text(self):
        """A bare URL (not markdown formatted) must NOT be converted to a hyperlink."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ CONTENT || p }}")
            template.save(str(template_path))

            store = SectionStore({
                "CONTENT": SectionRecord(
                    prompt="p",
                    output="See https://www.google.com for details.",
                    source="json",
                ),
            })
            self._process_docx(template_path, output_path, store=store)

            _, docxml = self._docx_xml(output_path)
            self.assertNotIn("<w:hyperlink", docxml)
            result = Document(str(output_path))
            joined = " ".join(p.text for p in result.paragraphs)
            self.assertIn("https://www.google.com", joined)

    def test_common_rich_markdown_features_render(self):
        """Footnotes, strike, sup/sub, and horizontal rules should survive DOCX conversion."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ CONTENT || p }}")
            template.save(str(template_path))

            store = SectionStore({
                "CONTENT": SectionRecord(
                    prompt="p",
                    output=(
                        "Term[^1] with ~~obsolete~~ note and H^^2^^O plus X~2~.\n\n"
                        "---\n\n"
                        "[^1]: Footnote body"
                    ),
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            texts = [p.text for p in result.paragraphs]
            joined = "\n".join(texts)

            self.assertIn("Footnote body", joined)
            self.assertIn("obsolete", joined)
            self.assertIn("H2O", joined)
            self.assertIn("X2", joined)
            self.assertIn("________________________________________", joined)

    def test_empty_placeholder_output_removes_line(self):
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("Before")
            template.add_paragraph("{{ MISSING }}")
            template.add_paragraph("After")
            template.save(str(template_path))

            self._process_docx(template_path, output_path)

            result = Document(str(output_path))
            non_empty = [paragraph.text.strip() for paragraph in result.paragraphs if paragraph.text.strip()]
            self.assertEqual(non_empty, ["Before", "After"])

    def test_multiple_placeholders_preserve_position(self):
        """Regression: content must appear at the position of its placeholder."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("Header")
            template.add_paragraph("{{ NAME || Name of system: }}")
            template.add_paragraph("")
            template.add_paragraph("{{ LICENSE }}")
            template.add_paragraph("")
            template.add_paragraph("Section heading")
            template.add_paragraph("{{ DESCRIPTION || Describe the system.")
            template.add_paragraph("Consider users and installation.")
            template.add_paragraph("}}")
            template.save(str(template_path))

            store = SectionStore({
                "NAME": SectionRecord(prompt="Name of system:", output="ACME Tool", source="json"),
                "LICENSE": SectionRecord(prompt="", output="MIT", source="json"),
                "DESCRIPTION": SectionRecord(
                    prompt="Describe the system.\nConsider users and installation.",
                    output="A great tool.",
                    source="json",
                ),
            })
            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            texts = [p.text for p in result.paragraphs]
            self.assertEqual(texts, ["Header", "ACME Tool", "", "MIT", "", "Section heading", "A great tool."])

    def testensure_required_styles_adds_missing_styles(self):
        """ensure_required_styles must add List Bullet / List Number when absent."""
        from docx.oxml.ns import qn

        doc = Document()
        # Remove List Bullet and List Number from the styles XML.
        styles_elem = doc.styles._element
        for style_el in styles_elem:
            name_el = style_el.find(qn("w:name"))
            if name_el is not None and name_el.get(qn("w:val")) in ("List Bullet", "List Number"):
                styles_elem.remove(style_el)

        existing_before = {s.name for s in doc.styles}
        self.assertNotIn("List Bullet", existing_before)
        self.assertNotIn("List Number", existing_before)

        ensure_required_styles(doc)

        existing_after = {s.name for s in doc.styles}
        self.assertIn("List Bullet", existing_after)
        self.assertIn("List Number", existing_after)

    def test_bulleted_list_output_does_not_crash(self):
        """A placeholder output containing a Markdown bullet list must not crash."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            # Strip list styles to simulate a sparse template.
            from docx.oxml.ns import qn as _qn
            styles_elem = template.styles._element
            for el in styles_elem:
                name_el = el.find(_qn("w:name"))
                if name_el is not None and name_el.get(_qn("w:val")) in ("List Bullet", "List Number"):
                    styles_elem.remove(el)
            template.add_paragraph("{{ ITEMS || List items }}")
            template.save(str(template_path))

            store = SectionStore({
                "ITEMS": SectionRecord(
                    prompt="List items",
                    output="- Alpha\n- Beta\n- Gamma",
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            joined = " ".join(p.text for p in result.paragraphs)
            self.assertIn("Alpha", joined)
            self.assertIn("Beta", joined)
            self.assertIn("Gamma", joined)


class NormalizeNewlinesTests(unittest.TestCase):
    """Tests for normalize_newlines edge cases."""

    def test_basic_newline_normalization_cases(self):
        cases = [
            ("line1\nline2", "line1  \nline2"),
            ("para1\n\npara2", "para1\n\npara2"),
            ("hello", "hello"),
            ("", ""),
        ]

        for original, expected in cases:
            with self.subTest(original=original):
                self.assertEqual(normalize_newlines(original), expected)

    def test_triple_newline_inserts_sentinel(self):
        """\\n\\n\\n should produce a sentinel for 1 extra blank paragraph."""
        result = normalize_newlines("a\n\n\nb")
        self.assertIn("<!-- EXTRA_NL:1 -->", result)

    def test_quad_newline_inserts_sentinel_2(self):
        """\\n\\n\\n\\n should produce a sentinel for 2 extra blank paragraphs."""
        result = normalize_newlines("a\n\n\n\nb")
        self.assertIn("<!-- EXTRA_NL:2 -->", result)

    def test_fenced_code_block_preserved(self):
        """Newlines inside fenced code blocks must NOT be altered."""
        code = "```\nline1\nline2\n```"
        result = normalize_newlines(code)
        self.assertEqual(result, code)

    def test_fenced_code_with_language(self):
        """Fenced code block with language specifier is preserved."""
        code = "```python\ndef foo():\n    pass\n```"
        result = normalize_newlines(code)
        self.assertEqual(result, code)

    def test_text_around_fenced_code(self):
        """Newlines outside code blocks are normalised; inside are not."""
        text = "before\nstuff\n```\na\nb\n```\nafter\nmore"
        result = normalize_newlines(text)
        self.assertIn("before  \nstuff", result)
        self.assertIn("```\na\nb\n```", result)
        self.assertIn("after  \nmore", result)

    def test_fence_variants_preserved(self):
        cases = [
            "```\nline1\nline2\n```",
            "```python\ndef foo():\n    pass\n```",
            "~~~\ncode\n~~~",
        ]

        for code in cases:
            with self.subTest(code=code):
                self.assertEqual(normalize_newlines(code), code)

    def test_newline_before_list_starter_upgraded_to_blank(self):
        cases = [
            ("prose\n- item1\n- item2", "- item1", "- item2"),
            ("intro\n1. first\n2. second", "1. first", "2. second"),
        ]

        for original, first_item, second_item in cases:
            with self.subTest(original=original):
                result = normalize_newlines(original)
                self.assertIn(f"\n\n{first_item}", result)
                self.assertNotIn(f"\n\n{second_item}", result)
                self.assertIn(f"\n{second_item}", result)

    def test_heading_before_bullet_list_stays_bare(self):
        """A heading line already closes a block; do not add extra blank before following list."""
        result = normalize_newlines("## Heading\n- item")
        self.assertNotIn("\n\n- item", result)
        self.assertIn("\n- item", result)

    def test_newline_before_heading_stays_bare(self):
        """# headings can interrupt paragraphs leave the \n bare (not hard break, not \n\n)."""
        result = normalize_newlines("text\n## Heading")
        self.assertNotIn("  \n#", result)
        self.assertNotIn("\n\n#", result)
        self.assertIn("\n## Heading", result)

    def test_newline_before_table_row_upgraded_to_blank(self):
        result = normalize_newlines("text\n| a | b |\n|---|---|")
        self.assertNotIn("  \n|", result)
        self.assertIn("\n\n| a | b |", result)

    def test_newline_before_blockquote_stays_bare(self):
        "> blockquotes can interrupt paragraphs leave the \n bare."
        result = normalize_newlines("text\n> quoted")
        self.assertNotIn("  \n>", result)
        self.assertNotIn("\n\n>", result)
        self.assertIn("\n> quoted", result)

    def test_bold_inline_not_mistaken_for_list(self):
        """**bold** starting a line is prose, not a list gets a hard break."""
        result = normalize_newlines("line1\n**bold** text")
        self.assertIn("  \n**bold**", result)


class PostprocessHtmlTests(unittest.TestCase):
    """Tests for postprocess_html behaviour."""

    def test_blank_p_between_consecutive_paragraphs(self):
        """An empty <p> should be inserted between adjacent <p> tags."""
        html = "<p>first</p>\n<p>second</p>"
        result = postprocess_html(html)
        self.assertIn("<p></p>", result)

    def test_no_blank_between_heading_and_paragraph(self):
        """A heading followed by a paragraph should NOT get a blank <p>."""
        html = "<h2>Title</h2>\n<p>Body</p>"
        result = postprocess_html(html)
        self.assertNotIn("<p></p>", result)

    def test_extra_nl_sentinel_expanded(self):
        """EXTRA_NL:2 sentinel should produce 2 blank <p> elements."""
        html = "<p>A</p>\n<!-- EXTRA_NL:2 -->\n<p>B</p>"
        result = postprocess_html(html)
        # Should have the sentinel expanded AND the consecutive-<p> blank.
        self.assertEqual(result.count("<p></p>"), 3) # 1 from adjacency + 2 from sentinel

    def test_blockquote_unwrapped(self):
        """<blockquote> paragraphs should get left-margin style."""
        html = "<blockquote><p>quoted</p></blockquote>"
        result = postprocess_html(html)
        self.assertNotIn("<blockquote>", result)
        self.assertIn("margin-left", result)

    def test_loose_list_item_paragraph_is_flattened(self):
        """A leading paragraph in a list item should be unwrapped before DOCX conversion."""
        html = "<ol><li><p><strong>Heading</strong>:</p><ul><li>Child</li></ul></li></ol>"
        result = postprocess_html(html)
        self.assertRegex(result, r"<p>1\.\s*<strong>\s*Heading</strong>:</p>")
        self.assertIn("<ul><li>Child</li></ul>", result)

    def test_single_paragraph_list_item_is_flattened(self):
        """A list item containing only one <p> must be flattened to avoid blank bullet line in DOCX."""
        html = "<ul><li><p><strong>Definition and Scope</strong>: Administrative logins refer</p></li></ul>"
        result = postprocess_html(html)
        self.assertIn("<ul><li><strong>Definition and Scope</strong>: Administrative logins refer</li></ul>", result)

    def test_unordered_list_is_converted_to_bulleted_paragraphs(self):
        """Unordered list HTML should remain list markup for DOCX list style handling."""
        html = "<ul><li>First</li><li>Second</li></ul>"
        result = postprocess_html(html)
        self.assertIn("<ul><li>First</li><li>Second</li></ul>", result)


class DocxNewlineOutputTests(_DocxPipelineMixin, unittest.TestCase):
    """End-to-end tests that verify newline handling produces correct DOCX output."""

    def test_double_newline_produces_blank_paragraph(self):
        """\\n\\n in content should produce a visible gap (empty paragraph) in DOCX."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ CONTENT || Provide content }}")
            template.save(str(template_path))

            store = SectionStore({
                "CONTENT": SectionRecord(
                    prompt="Provide content",
                    output="First paragraph\n\nSecond paragraph",
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            texts = [p.text for p in result.paragraphs]
            # There should be an empty paragraph between the two content paragraphs
            self.assertIn("First paragraph", texts)
            self.assertIn("Second paragraph", texts)
            first_idx = texts.index("First paragraph")
            second_idx = texts.index("Second paragraph")
            self.assertGreater(second_idx - first_idx, 1, "Expected blank paragraph between content")

    def test_triple_newline_produces_extra_blank(self):
        """\\n\\n\\n should produce more spacing than \\n\\n."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ CONTENT || Provide content }}")
            template.save(str(template_path))

            store = SectionStore({
                "CONTENT": SectionRecord(
                    prompt="Provide content",
                    output="Top\n\n\nBottom",
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            texts = [p.text for p in result.paragraphs]
            self.assertIn("Top", texts)
            self.assertIn("Bottom", texts)
            top_idx = texts.index("Top")
            bottom_idx = texts.index("Bottom")
            gap = bottom_idx - top_idx
            self.assertGreater(gap, 2, "Expected at least 2 blank paragraphs for \\n\\n\\n")

    def test_loose_numbered_item_keeps_text_on_same_line(self):
        """A numbered item followed by a blank line and nested bullets must keep its heading inline."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ CONTENT || Provide content }}")
            template.save(str(template_path))

            store = SectionStore({
                "CONTENT": SectionRecord(
                    prompt="Provide content",
                    output=(
                        "arrangement involved solely the act of copying and pasting.\n\n"
                        "2.  **Technical Implementation of Cookie Deletion**:\n\n"
                        "    *   **Target Scope**: The script"
                    ),
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            texts = [p.text for p in result.paragraphs]
            heading_line = next(t for t in texts if "Technical Implementation of Cookie Deletion:" in t)
            self.assertTrue(heading_line.startswith("2."))
            self.assertRegex(heading_line, r"^2\.\s+Technical")
            self.assertTrue(any("Target Scope: The script" in line for line in texts))

    def test_new_ordered_list_restarts_from_one(self):
        """A later ordered list in the same document should render from 1, not continue previous numbering."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("{{ FIRST || First list }}")
            template.add_paragraph("Between")
            template.add_paragraph("{{ SECOND || Second list }}")
            template.save(str(template_path))

            store = SectionStore({
                "FIRST": SectionRecord(prompt="First list", output="1. one\n2. two", source="json"),
                "SECOND": SectionRecord(
                    prompt="Second list",
                    output="1. **Code Generation and Authorship Methodology**: The source code was",
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            texts = [p.text for p in result.paragraphs if p.text.strip()]
            self.assertIn("1. one", texts)
            self.assertIn("2. two", texts)
            self.assertTrue(
                any(re.match(r"^1\.\s+Code Generation and Authorship Methodology: The source code was$", line) for line in texts)
            )

    def test_heading_followed_by_list_has_no_empty_gap(self):
        """A heading before a list should not get a converter-added blank paragraph."""
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"

            template = Document()
            template.add_paragraph("System administrator permissions", style="Heading 2")
            template.add_paragraph("{{ CONTENT || Provide content }}")
            template.save(str(template_path))

            store = SectionStore({
                "CONTENT": SectionRecord(
                    prompt="Provide content",
                    output="- **Definition and Scope**: Administrative",
                    source="json",
                ),
            })

            self._process_docx(template_path, output_path, store=store)

            result = Document(str(output_path))
            texts = [p.text for p in result.paragraphs]
            self.assertEqual(texts[0], "System administrator permissions")
            self.assertNotEqual(texts[1], "")
            self.assertIn("Definition and Scope: Administrative", texts[1])


if __name__ == "__main__":
    unittest.main()
